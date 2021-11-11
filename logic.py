from docx import Document
import functions
import importData



def mergeFile(filenameCO, filenameCD, filenameProgrmaOutput, lecturer, trimester, year):

    if((filenameCO == '') or filenameCD == ''):
        return -1

    try:
        document = Document(filenameCO)
        result = importData.setFileCD(filenameCD)
        if (result != 1):
            print("cant open file with data")
            return -1
        newdoc = document

        # --Updateing Header--start
        header = newdoc.sections[0].header
        runsHeader = header.paragraphs[0].runs
        runsHeader[0].text = importData.dataCourseCode[4]
        runsHeader[2].text = trimester + ', '
        runsHeader[3].text = str(year)
        runsHeader[4].text = runsHeader[5].text = runsHeader[6].text = ''
        # --Updateing Header--finished

        functions.replace_paragraph(newdoc, 3, importData.dataCourseCode[4])
        functions.replace_paragraph(newdoc, 4, importData.dataCourseTitle[4])
        functions.delete_paragraph(newdoc.paragraphs[5])
        functions.delete_paragraph(newdoc.paragraphs[6])
        functions.delete_paragraph(newdoc.paragraphs[6])

        #prereq = importData.dataPrerequisites[4]
        #current = functions.update_prerequisites(newdoc, prereq)
        #coreq = importData.dataCorequisites[4]
        #current = functions.update_corequisites(newdoc, coreq)

        print("Prerequisites OK") if functions.update_prerequisites(newdoc, importData.dataPrerequisites[4]) != -1 \
            else print("Prerequisites FAIL")
        print("Corequisites OK") if functions.update_corequisites(newdoc, importData.dataCorequisites[4]) != -1 \
            else print("Corequisites FAIL")
        current = functions.update_restrictions(newdoc, importData.dataRestrictions[4])
        print("Restrictions OK") if current != -1 else print("Restrictions FAIL")
        #functions.delete_paragraph(newdoc.paragraphs[current])

        functions.replace_paragraph(newdoc, current + 1, "NZQF Level: " + importData.dataNZQFLevel[4])
        functions.replace_paragraph(newdoc, current + 2, "Credits: " + importData.dataCredits[4])

        current = functions.update_lecturer(newdoc, lecturer)
        current = functions.update_lecturers(newdoc)
        current = functions.update_coordinator(newdoc)
        current = current - 5
        functions.delete_paragraph(newdoc.paragraphs[current])
        functions.delete_paragraph(newdoc.paragraphs[current])
        functions.delete_paragraph(newdoc.paragraphs[current])
        functions.update_courseAims(newdoc, importData.dataCourseAims[4])

        loList = importData.getLearningOutcomes()
        functions.update_learningOutcomes(newdoc, loList)

        importData.getAverageLearning()

        totalHours = importData.dataTotalHours
        functions.update_courseDuration(newdoc, totalHours)
        functions.update_learningHours(newdoc, totalHours)

        saList = importData.getSummativeAssesment()
        functions.update_courseAssesssments(newdoc, saList)

        newdoc.save(filenameProgrmaOutput)
        return 1;
    except:
        return -1;




