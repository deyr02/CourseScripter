from docx.shared import Pt

def find_paragraph(document, ptext):
    pnumber = 0
    for p in document.paragraphs:
        if (p.text == ptext):
            return pnumber
        pnumber = pnumber +1
    return -1

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def replace_paragraph(document, number, text):
 #   print(document.paragraphs[number].text)
    runs = document.paragraphs[number].runs
    for i in range(len(runs)):
        runs[i].text = ''
    document.paragraphs[number].runs[0].text = text
#    print(document.paragraphs[number].text)
    return 1

def update_prerequisites(document, prereq):
    firstline = find_paragraph(document, '+PREREQUISITES:+')
    seconLineline = find_paragraph(document, '+course code+ +title of the course+')
    if(firstline == -1 or (seconLineline - firstline != 1)): return -1

    replace_paragraph(document,firstline,'PREREQUISITES:')
    replace_paragraph(document,seconLineline,prereq)
    delete_paragraph(document.paragraphs[firstline + 2])
    delete_paragraph(document.paragraphs[firstline + 2])

    finishedParagraphNumber = firstline + 2
    return finishedParagraphNumber

def update_corequisites(document, coreq):
    firstline = find_paragraph(document, '+CO-REQUISITES:+')
    seconLineline = find_paragraph(document, '+course code + + title of the course+')
    if(firstline == -1 or (seconLineline - firstline != 1)): return -1

    replace_paragraph(document,firstline,'CO-REQUISITES:')
    replace_paragraph(document,seconLineline,coreq)
    delete_paragraph(document.paragraphs[firstline + 2])
    delete_paragraph(document.paragraphs[firstline + 2])

    finishedParagraphNumber = firstline + 2
    return finishedParagraphNumber

def update_restrictions(document, restrctions):
    firstline = find_paragraph(document, '+RESTRICTIONS:+')
    seconLineline = find_paragraph(document, '+course code + + title of the course+')
    if(firstline == -1 or (seconLineline - firstline != 1)): return -1

    replace_paragraph(document,firstline,'RESTRICTIONS:')
    replace_paragraph(document,seconLineline,restrctions)
    delete_paragraph(document.paragraphs[firstline + 2])
    delete_paragraph(document.paragraphs[firstline + 2])

    finishedParagraphNumber = firstline + 2
    return finishedParagraphNumber


def update_lecturer(document, lecturerInfo):
    lecturerLinesTemplate = ['Lecturer: 	',
                    'Room:	',
                    'Phone:	',
                    'Email: 	',
                    'Contact hour: 	']
    #lecturerInfo = ['Rakesh', '115', '+815-171718','rakesh@ais.ac.nz','every 4th of April, 10;15am']

    finishedParagraphNumber = -1
    firstline = -1
    i = 0
    for p in document.paragraphs:
        if(p.text == 'Lecturer: 	+Lecturer Name+' and document.paragraphs[i+21].text == 'Course AIM'):
            firstline = i;
        i = i+1
    if (firstline != -1):
        for j in range(firstline, firstline + len(lecturerLinesTemplate)):
            replace_paragraph(document,j,lecturerLinesTemplate[j-firstline] + lecturerInfo[j-firstline])
    else:
        return -1

    finishedParagraphNumber = firstline + len(lecturerLinesTemplate) + 1
    return finishedParagraphNumber

def update_lecturers(document):
    firstline = -1
    i = 0
    for p in document.paragraphs:
        if (p.text == 'Lecturers: 	+Lecturer Name 	email+' ):
            delete_paragraph(document.paragraphs[i])
            delete_paragraph(document.paragraphs[i])
            delete_paragraph(document.paragraphs[i])
            return i+1
        i = i + 1
    return -1

def update_coordinator(document):
    firstline = -1
    i = 0
    for p in document.paragraphs:
        if (p.text == 'Course Coordinator: 	+Lecturer Name+' and document.paragraphs[i + 7].text == 'Course AIM'):
            delete_paragraph(document.paragraphs[i])
            delete_paragraph(document.paragraphs[i])
            delete_paragraph(document.paragraphs[i])
            delete_paragraph(document.paragraphs[i])
            delete_paragraph(document.paragraphs[i])
            delete_paragraph(document.paragraphs[i])
            delete_paragraph(document.paragraphs[i])
            return i+1
        i = i + 1
    return -1

def update_courseAims(document, couseAims):
    firstline = find_paragraph(document, 'Course AIM')
    if(firstline == -1): return -1
    replace_paragraph(document,firstline+2,couseAims)

def update_courseDuration(document, hours):
    firstline = find_paragraph(document, 'Course DURATION')
    seconLineline = find_paragraph(document, '++FOR TRIMESTER-BASED COURSES++')
    if(firstline == -1 or (seconLineline - firstline != 1)): return -1
    replace_paragraph(document,seconLineline +1, 'The course will be conducted over a 14-week period with a typical teaching week consisting of '+ str(hours) +' hours of classes.')
    delete_paragraph(document.paragraphs[seconLineline])
# +2 The course will be conducted over a 14-week period with a typical teaching week consisting of +no of hours+ hours of classes.

def update_learningHours(document, totalHours):
    table = document.tables[0]
    if(table.rows[0].cells[0].text != 'Activity'):
        return -1
    for row in table.rows:
        if(row.cells[0].text == "TOTAL"):
            cell = row.cells[1]
            cell.paragraphs[0].runs[0].text = str(totalHours)
            cell.paragraphs[0].runs[1].text = ''

def update_learningOutcomes(document, loList):

    #find needed block (learning outcomes) and esure this is it (check second string as well)
    firstline = find_paragraph(document, 'Learning outcomes ')
    seconLineline = find_paragraph(document, 'The learners will be able to: ')
    if(firstline == -1 or (seconLineline - firstline != 1)): return -1
    pCopy = document.paragraphs[firstline + 2]  #string "Leaners will be able to:"
    #clear all the list under that line
    for i in range (firstline , firstline + 8):
        delete_paragraph(document.paragraphs[firstline+2])

    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    for i in range (1, len(loList)):
        p = pCopy;
        line = loList[i]
        p.text = line[1]
        par = document.paragraphs[firstline + 1 + i]
        p1 = par.insert_paragraph_before(str(line[0]) + ". " + line[1], style='Normal')     #print each learing outcome
    p1 = par.insert_paragraph_before("", style='Normal')  # add one empty line in the end of a block



def update_courseAssesssments(document, caList):
    firstline = find_paragraph(document, 'Course AssessmentS')
    seconLineline = find_paragraph(document, 'The assessments for this course are as follows: ')
    if(firstline == -1 or (seconLineline - firstline != 1)): return -1
    for i in range(0, len(caList)):
        str = caList[i]
    table = document.tables[3]
    if(table.rows[1].cells[0].text != '+ASSESSMENT 1+'):
        return -1

def deleteTellowLines(document):
    for p in document.paragraphs:
        if((len(p.text) > 2) and (p.text[0] == p.text[1] == '+' )):
            pasText = p.text
            pasNum = find_paragraph(document,pasText)
            if(pasNum != 30):
                delete_paragraph(document.paragraphs[pasNum])
