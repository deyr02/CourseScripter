from docx import *

#filename = 'TempleteCD(COMP503).docx'
filename = ''
document = ''

data1 = [0,0,0]
data2 = [0,0,1]
data3 = [0,1,1]
data4 = [0,1,0]

dataProgramme = [0, 0, 1, 'Programme', '']
dataCourseCode = [0,1,1, 'Course Code', '']
dataCourseTitle = [0,2,1, 'Course Title', '']
dataNZQFLevel = [0,3,1, 'NZQF Level', '']
dataCredits = [0,4,1,'Credits','']
dataPrerequisites = [0,5,1, 'Prerequisites','']
dataCorequisites = [0,6,1, 'Co-requisites', '']
dataRestrictions = [0,7,1,'Restrictions', '']
dataCourseAims = [0,8,1, 'Course Aims', '']
dataContent = [0,-2,1,'Content','']
dataDeliveryMethods = [0,-1,1,'Delivery Methods','']

dataTeachingHours = ''
dataSelfDirectedHours = ''
dataTotalHours = ''
dataTotalWeeks = ''


data_wrapper = [dataProgramme,
                dataCourseCode,
                dataCourseTitle,
                dataNZQFLevel,
                dataCredits,
                dataPrerequisites,
                dataCorequisites,
                dataRestrictions,
                dataCourseAims,
                dataContent]
'''
def setFileCD(fileName):
    global filename
    global document
    filename = fileName
    try:
        document = Document(filename)
    except:
        return -1

    al = getAverageLearning()
    sa = getSummativeAssesment()
    gd = getGeneralData()
    if((al == -1) or (sa == -1) or (gd != 1)): return 0
    else: return 1
'''
def setFileCD(fileName):
    global filename
    global document
    filename = fileName
    try:
        document = Document(filename)
        al = getAverageLearning()
        sa = getSummativeAssesment()
        gd = getGeneralData()
        lo = getLearningOutcomes()
        if ((al == -1) or (sa == -1) or (gd != 1)): return 0
        else:
            return 1
    except:
        return -1

def getGeneralData():
    for data in data_wrapper:
        table = document.tables[data[0]]
        row = table.rows[data[1]]
        cell = row.cells[data[2]]
        text = cell.text
        #print(text)
        #print(data[4])
        data[4] = text
    return 1

def getLearningOutcomes():
    lo = [0, '']
    loList = []
    for table in document.tables:
        counter = 0
        for row in table.rows:
            if row.cells[0].text == 'Learning\nOutcomes':
                lo = [counter, row.cells[2].text]
                loList.append(lo)
                counter = counter + 1
    return loList

def getSummativeAssesment():
    sa = ['', '', '']
    saList = []

    for table in document.tables:
        counter = 0
        for row in table.rows:
            #print(row.cells[0].text)
            if row.cells[0].text == 'Summative Assessment ':
                counter = counter + 1
                if(counter > 1):
                    #print(row.cells[3].text + " " + row.cells[4].text + " " + row.cells[5].text)
                    sa = [row.cells[3].text, row.cells[4].text, row.cells[5].text]
                    saList.append(sa)
    if(sa != ''): return saList
    else: return -1

def getAverageLearning():
    global dataTeachingHours
    global dataSelfDirectedHours
    global dataTotalHours
    global dataTotalWeeks

    for table in document.tables:
        counter = 0
        for row in table.rows:
            #print(row.cells[0].text)
            if row.cells[0].text == 'Average \nLearning ':
                counter = counter +1
                if(counter == 2):
                    #print(row.cells[2].text + " " + row.cells[3].text + " "+ row.cells[4].text + " "+ row.cells[5].text)
                    dataTeachingHours = row.cells[2].text
                    dataSelfDirectedHours = row.cells[3].text
                    dataTotalHours = row.cells[4].text
                    dataTotalWeeks = row.cells[5].text
                    return 1
    return -1

'''
getGeneralData()
getLearningOutcomes()
getSummativeAssesment()
getLearningOutcomes()
'''
'''
for table in document.tables:
    for row in table.rows:
        print (row.cells[0].text)
        if row.cells[0].text == 'Programme':
            print('fuck')
'''